"""Test C2/2.1 — trích ảnh kèm ngữ cảnh. Dựng .docx tổng hợp, chạy offline.

Mỗi test gắn với một hành vi ĐO ĐƯỢC trên 47 bản sizing thật (xem bảng số liệu ở
đầu `src/vision/anh.py`), không phải test cho vui.
"""
import struct
import tempfile
import zlib

from docx import Document
from docx.shared import Inches

from src.ingestion.docx_reader import read_docx
from src.vision.anh import (Anh, _alt_rac, _kich_thuoc_px, byte_anh, tach_caption,
                            trich_anh)


def _png(w: int, h: int, mau: bytes = b"\xff\x00\x00") -> bytes:
    """PNG hợp lệ nhỏ nhất có thể, để không phải phụ thuộc file mẫu."""
    raw = b"".join(b"\x00" + mau * w for _ in range(h))
    def khoi(ten: bytes, du_lieu: bytes) -> bytes:
        return (struct.pack(">I", len(du_lieu)) + ten + du_lieu
                + struct.pack(">I", zlib.crc32(ten + du_lieu) & 0xFFFFFFFF))
    ihdr = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)
    return (b"\x89PNG\r\n\x1a\n" + khoi(b"IHDR", ihdr)
            + khoi(b"IDAT", zlib.compress(raw)) + khoi(b"IEND", b""))


def _file_png(w=40, h=20) -> str:
    f = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    f.write(_png(w, h))
    f.close()
    return f.name


def _luu(doc) -> str:
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(f.name)
    f.close()
    return f.name


# --- caption ---------------------------------------------------------------
def test_nhan_ra_caption_kieu_hinh_so():
    assert tach_caption("Hình 3. Mô hình vật lý") == "Hình 3. Mô hình vật lý"
    assert tach_caption("Hình 1.2.5. Cấu hình CPU.")
    assert tach_caption("Figure 4 - Architecture")


def test_khong_nhan_cau_van_xuoi_dai_lam_caption():
    """"Hình" mở đầu một câu dài là văn xuôi, không phải caption."""
    dai = "Hình vẽ dưới đây mô tả " + "chi tiết " * 40
    assert tach_caption(dai) == ""


def test_dong_khong_bat_dau_bang_nhan_thi_khong_phai_caption():
    assert tach_caption("Tải CPU và RAM server 10.60.105.79") == ""


# --- alt text --------------------------------------------------------------
def test_alt_rac_dung_voi_cac_ca_do_duoc_tren_ho_so_that():
    # Bốn dạng này chiếm gần hết 154 alt text có nội dung trong 47 bản.
    for rac in ("IMG_256", "image13.png", "Picture 3", "",
                "cid:89122cbb-000c-4bab", r"C:\Users\ducdm\Pictures\before_upd.png"):
        assert _alt_rac(rac), rac
    assert not _alt_rac("Sơ đồ kiến trúc hệ thống BCCS3")


# --- kích thước ------------------------------------------------------------
def test_doc_kich_thuoc_png_va_gif():
    assert _kich_thuoc_px(_png(37, 21), "png") == (37, 21)
    gif = b"GIF89a" + struct.pack("<HH", 300, 120) + b"\x00" * 10
    assert _kich_thuoc_px(gif, "gif") == (300, 120)


def test_khong_doc_duoc_kich_thuoc_thi_tra_None_chu_khong_doan():
    """NT4: ảnh vector không có pixel — trả None, KHÔNG lấy 0 hay 1 làm mặc định."""
    assert _kich_thuoc_px(b"\x01\x00\x00\x00rac", "emf") == (None, None)
    assert _kich_thuoc_px(b"\x89PNG\r\n\x1a\n", "png") == (None, None)


# --- trích ảnh -------------------------------------------------------------
def test_trich_duoc_anh_kem_muc_va_caption_dung_sau():
    doc = Document()
    doc.add_paragraph("IV.1 Định cỡ module Postgres").bold = True
    doc.add_paragraph("Kết quả đo tải:")
    doc.add_picture(_file_png(), width=Inches(2))
    doc.add_paragraph("Hình 7. Tài nguyên khi có tải")

    d = read_docx(_luu(doc))
    kq = trich_anh(d)
    assert len(kq.anh) == 1
    a = kq.anh[0]
    assert a.dinh_dang == "png"
    assert (a.rong_px, a.cao_px) == (40, 20)
    assert a.caption == "Hình 7. Tài nguyên khi có tải"
    assert a.caption_nguon == "sau"          # đo thật: 76/107 caption nằm sau ảnh
    assert "Kết quả đo tải:" in a.truoc
    assert a.duong_dan_media.startswith("word/media/")
    assert byte_anh(d, a).startswith(b"\x89PNG")


def test_hai_anh_trong_mot_doan_ra_hai_ban_ghi():
    """C1 gộp cả đoạn thành một phần tử; 2.1 phải tách lại từng ảnh."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run().add_picture(_file_png(10, 10), width=Inches(1))
    p.add_run().add_picture(_file_png(20, 30), width=Inches(1))

    kq = trich_anh(read_docx(_luu(doc)))
    assert len(kq.anh) == 2
    assert [a.ma for a in kq.anh] == ["anh#0.0", "anh#0.1"]
    assert {(a.rong_px, a.cao_px) for a in kq.anh} == {(10, 10), (20, 30)}


def test_ngu_canh_gom_ca_muc_va_doan_hai_phia():
    doc = Document()
    doc.add_paragraph("Trước một")
    doc.add_paragraph("Trước hai")
    doc.add_picture(_file_png(), width=Inches(2))
    doc.add_paragraph("Sau một")

    a = trich_anh(read_docx(_luu(doc))).anh[0]
    assert a.truoc == ["Trước một", "Trước hai"]
    assert a.sau == ["Sau một"]
    assert "Trước hai" in a.ngu_canh() and "Sau một" in a.ngu_canh()


def test_khong_co_anh_thi_khong_co_canh_bao_thua():
    doc = Document()
    doc.add_paragraph("Chỉ có chữ")
    kq = trich_anh(read_docx(_luu(doc)))
    assert kq.anh == [] and kq.canh_bao == [] and kq.media_khong_dung == []


def test_ty_le_uu_tien_pixel_roi_moi_den_co_hien_thi():
    a = Anh(ma="x", element_index=0, rong_px=100, cao_px=50,
            emu_rong=999, emu_cao=1)
    assert a.ty_le == 2.0
    b = Anh(ma="x", element_index=0, emu_rong=300, emu_cao=100)
    assert b.ty_le == 3.0
    assert Anh(ma="x", element_index=0).ty_le is None


def test_location_giong_cach_nguoi_tham_dinh_trich_dan():
    a = Anh(ma="anh#3", element_index=3, section="IV.1.2", page=8)
    assert a.location == "Mục IV.1.2, trang 8"
    assert Anh(ma="anh#3", element_index=3).location == "phần tử #3"
