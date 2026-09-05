"""Test C1 — dựng .docx tổng hợp trong bộ nhớ, chạy offline.

Mỗi test ở đây gắn với một hành vi THẬT quan sát khi chạy C1 trên 48 bản sizing
lịch sử, nên chúng là hồi quy chứ không phải test cho vui.
"""
import io

import pytest
from docx import Document
from docx.enum.text import WD_BREAK

from src.ingestion.docx_reader import read_docx


def _save(doc) -> str:
    """Ghi ra file tạm và trả đường dẫn (read_docx nhận path)."""
    import tempfile
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(f.name)
    f.close()
    return f.name


def test_doc_rong_khong_lam_no():
    d = read_docx(_save(Document()))
    assert d.elements == []
    assert d.page_source == "none"


def test_bat_duoc_doan_bang_va_giu_thu_tu():
    doc = Document()
    doc.add_paragraph("Đoạn một")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "CPU"
    t.cell(0, 1).text = "143 Cint"
    t.cell(1, 0).text = "RAM"
    t.cell(1, 1).text = "176 GB"
    doc.add_paragraph("Đoạn hai")

    d = read_docx(_save(doc))
    kinds = [e.kind for e in d.elements]
    assert kinds == ["paragraph", "table", "paragraph"]
    assert [e.index for e in d.elements] == [0, 1, 2]
    tbl = d.tables()[0]
    assert tbl.rows[0] == ["CPU", "143 Cint"]
    assert "176 GB" in tbl.text


def test_so_muc_go_tay_la_ma_la_cap_chuong_a_rap_la_cap_duoi():
    """Bản THẬT gán 'Heading 1' cho cả mục con; con số mới cho biết cấp."""
    doc = Document()
    for t in ["I. THÔNG TIN HỆ THỐNG", "1. Thông tin chung",
              "2. Thông tin đầu vào", "II. GIẢI PHÁP", "1. Quan điểm"]:
        doc.add_paragraph(t, style="Heading 1")

    d = read_docx(_save(doc))
    secs = [e.section for e in d.elements if e.kind == "heading"]
    assert secs == ["I", "I.1", "I.2", "II", "II.1"]


def test_so_muc_nhieu_thanh_phan_duoc_ghep_voi_chuong():
    """Word hiện '1.5' ở cấp dưới; PNX trích 'Mục IV.1.5' -> phải ghép."""
    doc = Document()
    doc.add_paragraph("IV. TÍNH TOÁN CẤU HÌNH", style="Heading 1")
    doc.add_paragraph("1.5 Định cỡ Firewall", style="Heading 1")

    d = read_docx(_save(doc))
    assert [e.section for e in d.elements if e.kind == "heading"] == ["IV", "IV.1.5"]


def test_khong_co_so_la_ma_thi_so_a_rap_la_cap_chuong():
    doc = Document()
    doc.add_paragraph("1. Tổng quan", style="Heading 1")
    doc.add_paragraph("1.1 Chi tiết", style="Heading 1")

    d = read_docx(_save(doc))
    assert [e.section for e in d.elements if e.kind == "heading"] == ["1", "1.1"]


def test_moi_phan_tu_thua_ke_muc_dang_co_hieu_luc():
    doc = Document()
    doc.add_paragraph("III. TÍNH TOÁN", style="Heading 1")
    doc.add_paragraph("Nội dung thuộc chương III")

    d = read_docx(_save(doc))
    para = [e for e in d.elements if e.kind == "paragraph"][0]
    assert para.section == "III"
    assert para.section_title == "TÍNH TOÁN"


def test_by_section_lay_ca_muc_con():
    doc = Document()
    doc.add_paragraph("III. TÍNH TOÁN", style="Heading 1")
    doc.add_paragraph("1. Máy chủ", style="Heading 1")
    doc.add_paragraph("số liệu")
    doc.add_paragraph("IV. TỔNG HỢP", style="Heading 1")
    doc.add_paragraph("khác")

    d = read_docx(_save(doc))
    got = [e.text for e in d.by_section("III")]
    assert "số liệu" in got and "khác" not in got


def test_ngat_trang_thu_cong_tang_so_trang_va_canh_bao_uoc_luong():
    doc = Document()
    doc.add_paragraph("trang một")
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("trang hai")

    d = read_docx(_save(doc))
    assert d.page_source == "manual"
    assert d.elements[0].page == 1
    assert d.elements[-1].page == 2
    assert any("ƯỚC LƯỢNG" in w for w in d.warnings)


def test_khong_co_ngat_trang_thi_page_la_None_VA_NOI_RO():
    """NT4: không suy được thì nói không biết, không đoán bừa là trang 1."""
    doc = Document()
    doc.add_paragraph("một đoạn")

    d = read_docx(_save(doc))
    assert d.page_source == "none"
    assert all(e.page is None for e in d.elements)
    assert d.n_pages is None
    assert any("KHÔNG suy được số trang" in w for w in d.warnings)


def test_canh_bao_khi_khong_nhan_ra_de_muc_nao():
    doc = Document()
    doc.add_paragraph("chỉ có văn xuôi, không đề mục")
    d = read_docx(_save(doc))
    assert any("Không nhận ra đề mục" in w for w in d.warnings)


def test_location_ghep_muc_va_trang_theo_cach_nguoi_tham_dinh_viet():
    doc = Document()
    doc.add_paragraph("IV. TÍNH TOÁN", style="Heading 1")
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("nội dung")

    d = read_docx(_save(doc))
    last = d.elements[-1]
    assert last.location == "Mục IV, trang 2"


def test_location_khi_khong_co_gi_de_neo():
    doc = Document()
    doc.add_paragraph("văn xuôi")
    d = read_docx(_save(doc))
    assert d.elements[0].location == "phần tử #0"


def test_phan_tu_anh_giu_duoc_rid_de_lan_ra_file_anh():
    """C1 trước đây chỉ ghi "đoạn này có ảnh" — không đủ cho C2 đọc file nào.

    Hồi quy cho thay đổi 2026-09-04: `anh_refs` + `rels` là thứ để 2.1 tra ra
    `word/media/imageN.png`.
    """
    from docx.shared import Inches

    from tests.test_vision_anh import _file_png

    doc = Document()
    doc.add_picture(_file_png(30, 15), width=Inches(1.5))
    d = read_docx(_save(doc))

    anh = d.images()[0]
    assert anh.anh_refs and len(anh.anh_refs) == 1
    ref = anh.anh_refs[0]
    assert ref.rid and ref.neo == "inline"
    assert ref.emu_rong and ref.emu_cao          # cỡ hiển thị trong Word
    assert d.rels[ref.rid] in d.media
