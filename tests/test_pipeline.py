"""Test điều phối — chạy offline, không gọi model."""
import tempfile

from docx import Document

from src.extraction.schema import SizingCore
from src.ingestion.docx_reader import DocxDocument, Element, read_docx
from src.pipeline import canh_bao_nt4, chay


def _save(doc) -> str:
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(f.name)
    f.close()
    return f.name


def _doc(**kw) -> DocxDocument:
    return DocxDocument(path="giả.docx", **kw)


# ------------------------------------------------------------- NT4 --------
def test_anh_KHONG_bi_bo_qua_im_lang_du_giai_doan_1_chua_doc_anh():
    """767 ảnh trên 47 bản thật, và PNX liên tục nhận xét về ảnh sở cứ.
    "Bỏ qua" không được phép có nghĩa là im lặng (NT4)."""
    d = _doc(page_source="rendered", elements=[
        Element(index=0, kind="image", text="", page=4, section="III.2"),
        Element(index=1, kind="image", text="", page=5, section="III.3")])
    fs = canh_bao_nt4(d)
    anh = [f for f in fs if f.id == "NT4-ANH"]
    assert len(anh) == 1
    assert "2 hình ảnh" in anh[0].finding
    assert "Mục III.2, trang 4" in anh[0].computed_evidence
    assert anh[0].co_can_cu()               # NT2: căn cứ là con số code đếm


def test_canh_bao_anh_khong_do_ca_767_dong_ra_bao_cao():
    d = _doc(page_source="rendered", elements=[
        Element(index=i, kind="image", page=1, section="I") for i in range(30)])
    ev = canh_bao_nt4(d)[0].computed_evidence
    assert "30 ảnh" in ev and "25 ảnh khác" in ev


def test_khong_suy_duoc_so_trang_thi_noi_ra():
    d = _doc(page_source="none", elements=[Element(index=0, kind="paragraph", text="x")])
    assert any(f.id == "NT4-TRANG" for f in canh_bao_nt4(d))


def test_khong_co_anh_thi_khong_canh_bao_thua():
    d = _doc(page_source="rendered",
             elements=[Element(index=0, kind="paragraph", text="x", page=1)])
    assert canh_bao_nt4(d) == []


def test_canh_bao_cua_C1_duoc_chuyen_tiep_chu_khong_nuot():
    d = _doc(page_source="rendered", warnings=["không đọc được numbering.xml"],
             elements=[Element(index=0, kind="paragraph", text="x", page=1)])
    assert any("numbering.xml" in f.finding for f in canh_bao_nt4(d))


# ------------------------------------------------- mạch C1 -> C4 -> C7 ----
def test_chay_duoc_dau_cuoi_khong_can_model():
    """`bo_qua_trich_xuat` + `bo_qua_dinh_tinh` cho phép kiểm mạch mà không gọi mạng."""
    doc = Document()
    doc.add_paragraph("Tài liệu định cỡ thử.")
    kq = chay(_save(doc), bo_qua_trich_xuat=True, bo_qua_dinh_tinh=True)

    assert isinstance(kq.sizing, SizingCore)
    assert kq.ket_qua_dl and not kq.ket_qua_dt
    assert kq.thong_ke["c1_phan_tu"] == 1
    # Tài liệu rỗng -> toàn "thiếu thông tin", và mọi finding phải có căn cứ (NT2)
    assert all(f.co_can_cu() for f in kq.findings)
    bc = kq.bao_cao()
    assert "cố vấn" in bc.lower() or "advisory" in bc.lower()


def test_bao_cao_dung_duoc_tu_ket_qua_chay():
    doc = Document()
    doc.add_paragraph("x")
    kq = chay(_save(doc), bo_qua_trich_xuat=True, bo_qua_dinh_tinh=True)
    assert "Vòng 1" in kq.bao_cao() or "Vòng 2" in kq.bao_cao()
