"""Test 2.4 — cảnh báo NT4 về ảnh, tách theo loại. Chạy offline, không gọi model.

Trước 2.4 cảnh báo chỉ nói *"tài liệu có N hình ảnh chưa đọc được"*. Đúng, nhưng
người dùng không làm được gì với nó: 18 ảnh chụp dòng lệnh (nơi đặt số đo tải làm
sở cứ) và 1 sơ đồ đòi hai hành động hoàn toàn khác nhau.
"""
import struct
import tempfile
import zlib

import pytest
from docx import Document
from docx.shared import Inches

from src.ingestion.docx_reader import Element, read_docx
from src.pipeline import _canh_bao_anh, canh_bao_nt4
from src.vision import phan_loai as pl
from src.vision.phan_loai import NhomAnh, TomTatAnh, tom_tat_anh


def _png(rows) -> bytes:
    h, w = len(rows), len(rows[0])
    raw = b"".join(b"\x00" + b"".join(bytes(p) for p in row) for row in rows)

    def khoi(t, d):
        return (struct.pack(">I", len(d)) + t + d
                + struct.pack(">I", zlib.crc32(t + d) & 0xFFFFFFFF))
    return (b"\x89PNG\r\n\x1a\n"
            + khoi(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
            + khoi(b"IDAT", zlib.compress(raw)) + khoi(b"IEND", b""))


def _file(data: bytes) -> str:
    f = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    f.write(data)
    f.close()
    return f.name


# Nền đen + sọc sáng dày = mật độ chữ cao -> `console`. Nền trắng gần như không
# màu -> `anh_van_ban`. Hai ảnh này đo được ổn định, không nằm sát ngưỡng nào.
ANH_CONSOLE = [[(230, 230, 230) if x % 8 < 1 else (8, 8, 8) for x in range(400)]
               for _ in range(200)]
ANH_TRANG = [[(250, 250, 250) if not (10 < y < 14 and x % 60 < 40) else (40, 40, 40)
              for x in range(400)] for y in range(200)]


def _doc_co_anh(*anh_data):
    doc = Document()
    doc.add_paragraph("IV. TÍNH TOÁN CẤU HÌNH PHẦN CỨNG", style="Heading 1")
    for d in anh_data:
        doc.add_picture(_file(d), width=Inches(2))
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(f.name)
    f.close()
    return read_docx(f.name)


# --- tom_tat_anh -----------------------------------------------------------
def test_dem_dung_tung_loai_tren_tai_lieu_that():
    d = _doc_co_anh(_png(ANH_CONSOLE), _png(ANH_CONSOLE), _png(ANH_TRANG))
    tt = tom_tat_anh(d)
    assert tt.tong == 3 and tt.da_phan_loai
    theo = {n.loai: n.so_luong for n in tt.nhom}
    assert theo == {"console": 2, "anh_van_ban": 1}


def test_loai_nhieu_kha_nang_chua_so_do_dung_TRUOC():
    """Ảnh chụp dòng lệnh là nơi đặt sở cứ đo tải — người đọc cần thấy nó đầu tiên."""
    d = _doc_co_anh(_png(ANH_TRANG), _png(ANH_CONSOLE))
    assert [n.loai for n in tom_tat_anh(d).nhom] == ["console", "anh_van_ban"]


def test_tai_lieu_khong_anh_thi_tom_tat_rong():
    doc = Document()
    doc.add_paragraph("chỉ có chữ")
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(f.name)
    f.close()
    tt = tom_tat_anh(read_docx(f.name))
    assert tt.tong == 0 and tt.nhom == [] and not tt.da_phan_loai


# --- cảnh báo NT4 ----------------------------------------------------------
def test_moi_loai_ra_mot_canh_bao_rieng_kem_goi_y_rieng():
    d = _doc_co_anh(_png(ANH_CONSOLE), _png(ANH_TRANG))
    fs = _canh_bao_anh(d, d.images())
    ma = {f.id for f in fs}
    assert ma == {"NT4-ANH-1-CONSOLE", "NT4-ANH-2-ANH_VAN_BAN"}
    con = next(f for f in fs if "CONSOLE" in f.id)
    assert "dòng lệnh" in con.finding
    assert con.suggestion and con.suggestion != next(
        f for f in fs if "ANH_VAN_BAN" in f.id).suggestion


def test_canh_bao_van_la_muc_info_va_co_can_cu_dem_duoc():
    """NT2: căn cứ là con số code đếm. NT3: KHÔNG tự nâng mức độ ở đây."""
    d = _doc_co_anh(_png(ANH_CONSOLE))
    for f in _canh_bao_anh(d, d.images()):
        assert f.severity == "info"
        assert f.category == "khong_kiem_chung_duoc"
        assert f.co_can_cu() and "/1 ảnh" in f.computed_evidence


def test_khong_phan_loai_duoc_thi_LUI_VE_canh_bao_tong_kem_ly_do(monkeypatch):
    """NT4: mất tín hiệu pixel (thiếu Pillow, ảnh vector) thì vẫn phải nói ra tổng
    số ảnh, kèm lý do — không im lặng bỏ và cũng không đoán loại."""
    monkeypatch.setattr(
        "src.pipeline.tom_tat_anh",
        lambda doc: TomTatAnh(tong=2, nhom=[], canh_bao=["thiếu Pillow"],
                              da_phan_loai=False))
    els = [Element(index=0, kind="image", page=3, section="III.1"),
           Element(index=1, kind="image", page=4, section="III.2")]
    d = _doc_co_anh(_png(ANH_CONSOLE))
    fs = _canh_bao_anh(d, els)
    assert [f.id for f in fs] == ["NT4-ANH"]
    assert "2 hình ảnh" in fs[0].finding
    assert "thiếu Pillow" in fs[0].computed_evidence


def test_khong_liet_ke_het_vi_tri_khi_qua_nhieu_anh(monkeypatch):
    monkeypatch.setattr(
        "src.pipeline.tom_tat_anh",
        lambda doc: TomTatAnh(
            tong=30, da_phan_loai=True,
            nhom=[NhomAnh(loai="console", so_luong=30,
                          vi_tri=[f"trang {i}" for i in range(1, 6)])]))
    d = _doc_co_anh(_png(ANH_CONSOLE))
    ev = _canh_bao_anh(d, d.images())[0].computed_evidence
    assert "30/30 ảnh" in ev and "25 ảnh khác" in ev


def test_loai_la_khong_co_nhan_trong_config_van_khong_lam_vo_bao_cao(monkeypatch):
    monkeypatch.setattr(
        "src.pipeline.tom_tat_anh",
        lambda doc: TomTatAnh(tong=1, da_phan_loai=True,
                              nhom=[NhomAnh(loai="loai_moi", so_luong=1,
                                            vi_tri=["trang 1"])]))
    d = _doc_co_anh(_png(ANH_CONSOLE))
    f = _canh_bao_anh(d, d.images())[0]
    assert f.id == "NT4-ANH-1-LOAI_MOI" and "loai_moi" in f.finding


def test_canh_bao_anh_nam_trong_canh_bao_nt4_chung():
    d = _doc_co_anh(_png(ANH_CONSOLE))
    assert any(f.id.startswith("NT4-ANH") for f in canh_bao_nt4(d))


@pytest.mark.parametrize("loai", ["console", "dashboard", "anh_van_ban", "so_do",
                                  "chua_ro"])
def test_moi_loai_deu_co_nhan_va_goi_y_trong_config(loai):
    """Nhãn hiển thị là DỮ LIỆU (`config/report_labels.yaml`), không hard-code."""
    from src.reporting.report import load_labels
    nh = load_labels().anh_loai
    assert loai in nh and nh[loai].get("ten") and nh[loai].get("goi_y")


def test_thu_tu_uu_tien_khop_voi_danh_sach_loai_cua_2_2():
    """Đổi tên một loại ở 2.2 mà quên sửa thứ tự ưu tiên thì loại đó tụt xuống cuối
    một cách âm thầm — test này chặn đúng chỗ đó."""
    from typing import get_args
    assert set(get_args(pl.Loai)) == {"so_do", "console", "dashboard",
                                      "anh_van_ban", "chua_ro"}


def test_thu_tu_uu_tien_song_qua_C7_chu_khong_bi_xep_lai_theo_bang_chu_cai():
    """Hồi quy cho một lỗi LIÊN THÀNH PHẦN đã gặp thật: C7 xếp finding cùng mức độ
    theo `id`, nên `chua_ro` từng hiện TRƯỚC `console` trong báo cáo bản Vtag —
    đúng thứ tự ưu tiên dựng ở 2.2 bị vứt đi."""
    from src.reporting.report import build_report

    d = _doc_co_anh(_png(ANH_CONSOLE), _png(ANH_TRANG))
    bc = build_report(_canh_bao_anh(d, d.images()))
    assert bc.index("dòng lệnh") < bc.index("chụp bảng hoặc văn bản")
