"""Test C2/2.3 — đọc ảnh bằng vision. Chạy OFFLINE bằng client giả.

Chạy thật cần model, nhưng mọi thứ quyết định chất lượng — cổng chống bịa, việc
CODE ra số, chọn loại ảnh, xuống cấp khi không đọc được — đều test được ở đây.
"""
import struct
import tempfile
import zlib

import pytest
from docx import Document
from docx.shared import Inches

from src.ingestion.docx_reader import read_docx
from src.llm.client import ExtractionFailed
from src.vision import doc_anh as da
from src.vision.anh import Anh
from src.vision.doc_anh import (LOAI_MAC_DINH, DocAnh, DocConsole, DocSoDo,
                                SoLieuAnh, dong_goi, neo_duoc, thanh_finding,
                                uoc_tinh_luot_goi_anh)
from src.vision.phan_loai import co_pillow

can_pillow = pytest.mark.skipif(
    not co_pillow(),
    reason="cần Pillow để đo đặc trưng ảnh (C2 mục 2.2) — cài bằng `uv sync`")


# --- hạ tầng giả -----------------------------------------------------------
class ClientGia:
    """Trả sẵn kết quả và GHI LẠI thông điệp, để kiểm cách đóng gói lời gọi."""

    def __init__(self, tra_ve):
        self.tra_ve = list(tra_ve)
        self.messages = []
        self.vision_model = "fake-vision"

    def extract(self, schema, messages, *, model=None, **kw):
        self.messages.append(messages)
        r = self.tra_ve.pop(0)
        if isinstance(r, Exception):
            raise r
        return r


def _png(rows) -> bytes:
    h, w = len(rows), len(rows[0])
    raw = b"".join(b"\x00" + b"".join(bytes(p) for p in row) for row in rows)

    def khoi(t, d):
        return (struct.pack(">I", len(d)) + t + d
                + struct.pack(">I", zlib.crc32(t + d) & 0xFFFFFFFF))
    return (b"\x89PNG\r\n\x1a\n"
            + khoi(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
            + khoi(b"IDAT", zlib.compress(raw)) + khoi(b"IEND", b""))


ANH_CONSOLE = [[(230, 230, 230) if x % 8 < 1 else (8, 8, 8) for x in range(400)]
               for _ in range(200)]
ANH_TRANG = [[(250, 250, 250) if not (10 < y < 14 and x % 60 < 40) else (40, 40, 40)
              for x in range(400)] for y in range(200)]


def _anh(**kw) -> Anh:
    kw.setdefault("ma", "anh#1")
    kw.setdefault("element_index", 1)
    kw.setdefault("dinh_dang", "png")
    return Anh(**kw)


def _so(nhan, raw, trich, don_vi=""):
    return SoLieuAnh(nhan=nhan, gia_tri_raw=raw, trich_dan=trich, don_vi=don_vi)


# --- mặc định do người dùng chốt -------------------------------------------
def test_mac_dinh_chi_doc_so_do_va_console():
    """Chốt 2026-09-05. Đổi mặc định là đổi chi phí lượt chạy — phải hỏi lại."""
    assert LOAI_MAC_DINH == ("so_do", "console")


# --- NT2: cổng chống bịa ---------------------------------------------------
def test_neo_duoc_doi_gia_tri_phai_nam_trong_chinh_trich_dan():
    assert neo_duoc("65808076", "KiB Mem : 65808076 total, 3294864 free")
    assert not neo_duoc("128957", "KiB Mem : 65808076 total")
    assert not neo_duoc("", "KiB Mem : 65808076 total")
    assert not neo_duoc("65808076", "")


def test_gia_tri_khong_nam_trong_trich_dan_thi_BI_LOAI():
    c = ClientGia([DocConsole(doc_duoc=True, so_lieu=[
        _so("RAM", "65808076", "KiB Mem : 65808076 total"),
        _so("CPU", "99", "Tasks: 368 total"),          # 99 không có trong trích dẫn
    ])])
    kq = DocAnh(c).doc_mot(_anh(), "console", _png(ANH_CONSOLE))
    assert [s.nhan for s in kq.so_lieu] == ["RAM"]
    assert kq.bo_vi_khong_neo == 1


def test_moi_gia_tri_deu_bi_loai_thi_coi_nhu_KHONG_doc_duoc():
    """Giữ lại một kết quả rỗng mà vẫn nói "đọc được" là nói dối bằng im lặng."""
    c = ClientGia([DocConsole(doc_duoc=True, so_lieu=[
        _so("RAM", "111", "không có số nào ở đây")])])
    d = DocAnh(c)
    kq = d.doc_mot(_anh(), "console", _png(ANH_CONSOLE))
    assert not kq.doc_duoc and "không nằm trong trích dẫn" in kq.ly_do
    assert d.tk.doc_duoc == 0 and d.tk.khong_doc_duoc == 1


# --- NT1: CODE ra số, không phải model ------------------------------------
def test_luoc_do_KHONG_co_truong_so_hoc_nao():
    """Model chỉ được trả chuỗi. Có một trường float là mở đường cho nó tự tính."""
    for ten, f in SoLieuAnh.model_fields.items():
        assert f.annotation is str, f"{ten} phải là chuỗi, không phải {f.annotation}"


def test_so_luong_nghia_giu_ca_hai_cach_doc():
    """"1.500" là 1500 hay 1,5 — giữ cả hai, y như thang suy luận của 1.4."""
    c = ClientGia([DocConsole(doc_duoc=True, so_lieu=[
        _so("RAM", "1.500", "RAM: 1.500 GB", "GB")])])
    s = DocAnh(c).doc_mot(_anh(), "console", _png(ANH_CONSOLE)).so_lieu[0]
    assert s.gia_tri == 1500.0 and s.luong_nghia and s.gia_tri_khac == 1.5
    assert s.raw == "1.500"          # nguyên văn vẫn giữ, cho NT2


# --- NT4: xuống cấp có kiểm soát -------------------------------------------
def test_model_bao_khong_doc_duoc_thi_ghi_lai_ly_do_chu_khong_bia():
    c = ClientGia([DocConsole(doc_duoc=False, ly_do_khong_doc_duoc="ảnh mờ")])
    kq = DocAnh(c).doc_mot(_anh(), "console", _png(ANH_CONSOLE))
    assert not kq.doc_duoc and kq.ly_do == "ảnh mờ" and kq.so_lieu == []


def test_so_do_bao_doc_duoc_nhung_rong_thi_van_la_khong_doc_duoc():
    c = ClientGia([DocSoDo(doc_duoc=True)])
    kq = DocAnh(c).doc_mot(_anh(), "so_do", _png(ANH_TRANG))
    assert not kq.doc_duoc and "không nêu thành phần" in kq.ly_do


def test_loi_goi_mo_hinh_khong_lam_dung_ca_luot_chay():
    c = ClientGia([ExtractionFailed(3, "hỏng", "")])
    d = DocAnh(c)
    kq = d.doc_mot(_anh(), "console", _png(ANH_CONSOLE))
    assert not kq.doc_duoc and "lỗi gọi mô hình" in kq.ly_do
    assert d.tk.luot_goi_hong == 1 and d.tk.loi


def test_dinh_dang_gateway_khong_nhan_thi_BO_TRUOC_khi_goi():
    """emf/wmf gọi đi là đốt ~40 giây để nhận về một lỗi."""
    c = ClientGia([])
    d = DocAnh(c)
    kq = d.doc_mot(_anh(dinh_dang="emf"), "so_do", b"\x01\x00\x00\x00rac")
    assert not kq.doc_duoc and "emf" in kq.ly_do
    assert c.messages == [] and d.tk.bo_qua_dinh_dang == 1


def test_anh_qua_lon_thi_bo_kem_ly_do(monkeypatch):
    monkeypatch.setattr(da, "MAX_BYTE", 10)
    monkeypatch.setattr(da, "thu_nho", lambda data, dd, **kw: (data, dd))
    c = ClientGia([])
    d = DocAnh(c)
    kq = d.doc_mot(_anh(), "console", _png(ANH_CONSOLE))
    assert not kq.doc_duoc and "vượt mức gửi được" in kq.ly_do
    assert c.messages == [] and d.tk.bo_qua_qua_lon == 1


# --- đóng gói lời gọi ------------------------------------------------------
def test_thong_diep_co_anh_dang_data_url_va_nhac_rieng_theo_loai():
    c = ClientGia([DocConsole(doc_duoc=True, so_lieu=[
        _so("RAM", "16", "Mem: 16 GB")])])
    DocAnh(c).doc_mot(_anh(), "console", _png(ANH_CONSOLE))
    tin = c.messages[0]
    assert tin[0]["role"] == "system" and "KHÔNG tính toán" in tin[0]["content"]
    phan = {p["type"]: p for p in tin[1]["content"]}
    assert phan["image_url"]["image_url"]["url"].startswith("data:image/png;base64,")
    assert "dòng lệnh" in phan["text"]["text"]


def test_so_do_dung_loi_nhac_khac_console():
    c = ClientGia([DocSoDo(doc_duoc=True, thanh_phan=["Kafka"])])
    DocAnh(c).doc_mot(_anh(), "so_do", _png(ANH_TRANG))
    txt = [p for p in c.messages[0][1]["content"] if p["type"] == "text"][0]["text"]
    assert "sơ đồ" in txt and "dòng lệnh" not in txt


def test_ngu_canh_van_ban_duoc_gan_nhan_la_KHONG_duoc_lay_so_tu_do():
    """Chữ quanh ảnh giúp model hiểu bối cảnh, nhưng số phải đọc TỪ ẢNH."""
    c = ClientGia([DocConsole(doc_duoc=True, so_lieu=[
        _so("RAM", "16", "Mem: 16 GB")])])
    DocAnh(c).doc_mot(_anh(truoc=["Tải RAM: 999 GB"]), "console", _png(ANH_CONSOLE))
    txt = [p for p in c.messages[0][1]["content"] if p["type"] == "text"][0]["text"]
    assert "KHÔNG được lấy số từ đây" in txt


def test_dong_goi_tra_None_khi_dinh_dang_khong_ho_tro_duoc_thu_nho(monkeypatch):
    monkeypatch.setattr(da, "MAX_BYTE", 5)
    monkeypatch.setattr(da, "thu_nho", lambda data, dd, **kw: (data, dd))
    assert dong_goi(b"0123456789", "png") is None


# --- chọn ảnh --------------------------------------------------------------
def _doc_co_anh(*anh_data):
    doc = Document()
    doc.add_paragraph("IV. TÍNH TOÁN", style="Heading 1")
    for d in anh_data:
        f = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        f.write(d)
        f.close()
        doc.add_picture(f.name, width=Inches(2))
    g = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(g.name)
    g.close()
    return read_docx(g.name)


@can_pillow
def test_chi_goi_model_cho_loai_da_chon():
    """Ảnh chụp bảng không nằm trong mặc định thì KHÔNG được đốt một lượt gọi."""
    d = _doc_co_anh(_png(ANH_CONSOLE), _png(ANH_TRANG), _png(ANH_CONSOLE))
    c = ClientGia([DocConsole(doc_duoc=True, so_lieu=[_so("a", "1", "a 1")]),
                   DocConsole(doc_duoc=True, so_lieu=[_so("b", "2", "b 2")])])
    kq = DocAnh(c).run(d)
    assert len(c.messages) == 2 and len(kq) == 2
    assert all(k.loai == "console" for k in kq)


@can_pillow
def test_bat_them_loai_bang_tham_so():
    d = _doc_co_anh(_png(ANH_CONSOLE), _png(ANH_TRANG))
    c = ClientGia([DocConsole(doc_duoc=True, so_lieu=[_so("a", "1", "a 1")])] * 2)
    kq = DocAnh(c, loai=("console", "anh_van_ban")).run(d)
    assert len(kq) == 2


@can_pillow
def test_uoc_tinh_dem_truoc_khi_tieu_tien():
    d = _doc_co_anh(_png(ANH_CONSOLE), _png(ANH_CONSOLE), _png(ANH_TRANG))
    u = uoc_tinh_luot_goi_anh(d)
    assert u["tong_anh"] == 3 and u["se_doc"] == 2
    assert u["theo_loai"]["console"] == 2


def test_tai_lieu_khong_anh_thi_khong_goi_lan_nao():
    doc = Document()
    doc.add_paragraph("chỉ có chữ")
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(f.name)
    f.close()
    c = ClientGia([])
    assert DocAnh(c).run(read_docx(f.name)) == [] and c.messages == []


# --- finding ---------------------------------------------------------------
def test_finding_luon_co_can_cu_du_doc_duoc_hay_khong():
    c = ClientGia([DocConsole(doc_duoc=True, so_lieu=[_so("RAM", "16", "Mem: 16 GB")]),
                   DocConsole(doc_duoc=False, ly_do_khong_doc_duoc="ảnh mờ")])
    d = DocAnh(c)
    tot = thanh_finding(d.doc_mot(_anh(), "console", _png(ANH_CONSOLE)))
    xau = thanh_finding(d.doc_mot(_anh(ma="anh#2"), "console", _png(ANH_CONSOLE)))
    assert tot.co_can_cu() and xau.co_can_cu()
    assert tot.confidence == "vua"        # model đọc, người phải kiểm lại
    assert "ảnh mờ" in xau.finding


@pytest.mark.parametrize("loai", ["so_do", "console"])
def test_moi_loai_mac_dinh_deu_co_luoc_do_va_loi_nhac(loai):
    assert loai in da.LUOC_DO and loai in da.NHAC


# --- nối vào pipeline: mặc định TẮT ---------------------------------------
def test_pipeline_KHONG_goi_vision_khi_chua_bat():
    """Bảo vệ quyết định 2026-09-05 bằng code: lượt đo recall (B1) chạy SẠCH,
    không kèm 2.3. Nếu ai đó đổi mặc định thành True, test này phải đỏ."""
    import inspect

    from src.pipeline import chay
    tham_so = inspect.signature(chay).parameters
    assert tham_so["doc_anh"].default is False


def test_pipeline_bat_len_thi_finding_anh_vao_bao_cao(monkeypatch):
    from src.pipeline import chay
    from src.vision.doc_anh import KetQuaDocAnh

    class C2Gia:
        def __init__(self, *a, **kw):
            self.tk = type("T", (), {"__dict__": {"luot_goi": 1}})()

        def run(self, doc):
            return [KetQuaDocAnh(ma_anh="anh#1", loai="console", location="trang 1",
                                 doc_duoc=False, ly_do="ảnh mờ")]

    monkeypatch.setattr("src.vision.doc_anh.DocAnh", C2Gia)
    d = _doc_co_anh(_png(ANH_CONSOLE))
    kq = chay(d.path, client=object(), bo_qua_trich_xuat=True,
              bo_qua_dinh_tinh=True, doc_anh=True)
    assert any(f.id == "C2-ANH-anh#1" for f in kq.findings)
    assert kq.thong_ke["c2"]["luot_goi"] == 1
