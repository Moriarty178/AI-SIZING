"""1.16 — sinh MẪU WORD chuẩn thẳng từ 57 mục checklist thẩm định.

Trước đây mục này chờ Phụ lục 01 của Guideline. Không cần nữa: **checklist thẩm định
vốn đã là danh mục đề mục bắt buộc**, đúng thứ tự và phân cấp — chính là thứ Phụ lục 01
lẽ ra cung cấp. Không cần LLM, không phụ thuộc thành phần nào.

Đây là cách rẻ nhất để hạ **rủi ro R4** (định dạng tài liệu lộn xộn): đỡ người viết ngay
từ khâu tạo, thay vì bắt lỗi sau khi họ viết xong. Một mục viết đúng chỗ ngay từ đầu thì
C1 đọc được, C3 trích được, và người thẩm định chấm checklist theo cùng một mạch.

Nguồn là `docs/rules/.tmp-checklist/.md` — bản số hoá thô của file Excel gốc, **dữ liệu
chứ không phải code** (NT3). Tiêu chí đạt của từng mục được chép **nguyên văn** vào mẫu
làm lời nhắc, nên người viết thấy đúng câu người thẩm định sẽ dùng để chấm.

**Khối 20 mục của phần III lặp cho MỌI phân hệ**, không chỉ Application và Database —
checklist gốc chỉ liệt kê hai cái đó làm ví dụ. `phan_he_them` sinh thêm bản sao cho
Redis, Kafka, K8s… đúng số phân hệ thật của hệ thống.
"""
from __future__ import annotations

import re
from dataclasses import dataclass

DUONG_DAN_CHECKLIST = "docs/rules/.tmp-checklist/items.md"

# Ba lỗi trong file Excel gốc, đã chốt cách xử lý ở mục 0.1 và ghi trong
# `docs/rules/rules-checklist-flat.md`: KHÔNG sửa file nguồn, chỉ vá lúc đọc.
#   ô A42 ghi `3.1.2` nhưng theo ngữ cảnh là `3.2` (tiêu đề khối Database)
#   dòng 18 và dòng 50 thiếu số thứ tự  ->  `2.10a` và `3.2.7a`
SUA_LOI_NGUON = {"42": "3.2", "18": "2.10a", "50": "3.2.7a"}

_HANG = re.compile(r"^\|")
_DAM = re.compile(r"\*\*(.*?)\*\*", re.DOTALL)


@dataclass
class MucChecklist:
    dong: str
    tt: str
    hang_muc: str
    tieu_chi: str = ""
    tham_chieu: str = ""

    @property
    def la_chuong(self) -> bool:
        return self.tt in ("I", "II", "III")

    @property
    def la_tieu_de_khoi(self) -> bool:
        """`3.1` / `3.2` — tiêu đề khối phân hệ, không phải mục để điền."""
        return bool(re.fullmatch(r"\d+\.\d+", self.tt)) and self.tt.startswith("3.")

    @property
    def cap(self) -> int:
        if self.tt == "A":
            return 0
        if self.la_chuong:
            return 1
        return 2 if self.la_tieu_de_khoi else 3


def _o(hang: str) -> list[str]:
    return [c.strip() for c in hang.strip().strip("|").split("|")]


def doc_checklist(duong_dan: str = DUONG_DAN_CHECKLIST) -> list[MucChecklist]:
    """Đọc bảng checklist thành danh sách mục. Bỏ đậm, vá ba lỗi nguồn đã biết."""
    ra: list[MucChecklist] = []
    with open(duong_dan, encoding="utf-8") as f:
        hang = [l for l in f if _HANG.match(l)]
    for h in hang[2:]:                       # bỏ dòng tiêu đề + dòng gạch
        c = _o(h)
        if len(c) < 3:
            continue
        dong, tt = c[0], _DAM.sub(r"\1", c[1]).strip()
        if tt in ("_(sót)_", ""):
            tt = SUA_LOI_NGUON.get(dong, "")
        else:
            tt = SUA_LOI_NGUON.get(dong, tt)
        ra.append(MucChecklist(
            dong=dong, tt=tt,
            hang_muc=_DAM.sub(r"\1", c[2]).strip(),
            tieu_chi=_DAM.sub(r"\1", c[3]).strip() if len(c) > 3 else "",
            tham_chieu=_DAM.sub(r"\1", c[4]).strip() if len(c) > 4 else ""))
    return ra


def khoi_phan_he(mucs: list[MucChecklist], tien_to: str = "3.1") -> list[MucChecklist]:
    """20 mục dùng chung của một phân hệ, lấy khối Application làm bản chuẩn."""
    return [m for m in mucs if m.tt.startswith(tien_to + ".")]


# ---------------------------------------------------------------------------
def dung_mau(mucs: list[MucChecklist], *, phan_he_them: list[str] | None = None,
             ten_he_thong: str = "<TÊN HỆ THỐNG>"):
    """Dựng `docx.Document` theo cấu trúc checklist. Trả về Document chưa lưu."""
    from docx import Document
    from docx.shared import Pt

    d = Document()
    tieu_de = next((m for m in mucs if m.tt == "A"), None)
    d.add_heading(f"TÍNH TOÁN ĐỊNH CỠ HỆ THỐNG {ten_he_thong}", level=0)

    p = d.add_paragraph()
    r = p.add_run(
        "Mẫu này sinh tự động từ checklist thẩm định của đơn vị thẩm định "
        f"({tieu_de.hang_muc if tieu_de else 'checklist'}). Mỗi đề mục kèm sẵn "
        "tiêu chí đạt — chép nguyên văn từ checklist — để biết cần viết gì. "
        "Xoá dòng in nghiêng sau khi đã điền.")
    r.italic = True
    r.font.size = Pt(9)

    def them_muc(m: MucChecklist, cap: int) -> None:
        so = f"{m.tt}. " if m.tt else ""
        d.add_heading(f"{so}{m.hang_muc}", level=min(cap, 4))
        if m.tieu_chi:
            hp = d.add_paragraph()
            hr = hp.add_run(f"Tiêu chí đạt: {m.tieu_chi}")
            hr.italic = True
            hr.font.size = Pt(9)
        if m.tham_chieu:
            tp = d.add_paragraph()
            tr = tp.add_run(f"Tham chiếu: {m.tham_chieu}")
            tr.italic = True
            tr.font.size = Pt(8)
        d.add_paragraph("")                  # chỗ trống để người viết điền

    trong_phan_iii = False
    for m in mucs:
        if m.tt == "A":
            continue
        if m.la_chuong:
            d.add_heading(f"{m.tt}. {m.hang_muc}", level=1)
            trong_phan_iii = m.tt == "III"
            continue
        if m.la_tieu_de_khoi:
            d.add_heading(f"{m.tt}. {m.hang_muc}", level=2)
            continue
        them_muc(m, 2 if not trong_phan_iii else 3)

    # Khối 20 mục lặp cho mọi phân hệ khác — checklist gốc chỉ nêu App và DB.
    for ten in phan_he_them or []:
        d.add_heading(f"Phân hệ {ten}", level=2)
        for m in khoi_phan_he(mucs):
            m2 = MucChecklist(dong=m.dong, tt="", hang_muc=m.hang_muc,
                              tieu_chi=m.tieu_chi, tham_chieu=m.tham_chieu)
            them_muc(m2, 3)
    return d
