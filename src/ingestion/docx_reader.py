"""C1 — đọc .docx thành danh sách phần tử CÓ VỊ TRÍ.

Vì sao phải giữ vị trí: mọi finding đều phải neo được vào chỗ trong tài liệu
(schema C7 có trường `location`), và nhãn vàng lấy từ PNX neo theo **"Trang N"**
và **"Mục IV.1.2"** (xem `docs/0.7-nhan-vang-tu-pnx.md`). Không có hai thứ đó thì
không đối chiếu được finding với nhãn, tức không đo được recall.

Ba khó khăn của tài liệu THẬT (đã chốt ở 0.11: đọc bản người dùng tự viết, không
phải bản web app xuất ra):

1. **Ít khi dùng Heading style.** Người viết đánh số tay ("IV.1.2 Định cỡ máy chủ")
   và bôi đậm. Nên nhận diện đề mục bằng CẢ style LẪN mẫu đánh số, không chỉ style.
2. **Số trang không có sẵn trong .docx.** Chỉ suy được từ dấu ngắt trang Word ghi
   lại khi kết xuất (`w:lastRenderedPageBreak`) và ngắt trang thủ công. Có file
   không có dấu nào — khi đó `page` là `None` và phải NÓI RÕ, không đoán bừa (NT4).
3. **Bảng lồng bảng, ảnh xen giữa.** Ảnh không bỏ qua im lặng: ghi lại thành phần
   tử `image` kèm vị trí để C2 (GĐ 2) xử lý và để NT4 sinh cảnh báo.
"""
from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass, field

from docx import Document as _Docx
from docx.oxml.ns import qn

from .numbering import load_numbering, para_num_ref

# "IV.1.2 ...", "4.1.2. ...", "Mục 3.2 ..." — số đề mục người viết gõ tay
_NUM = re.compile(
    r"^\s*(?:mục\s+)?"
    r"(?P<num>(?:[IVXLCDM]+|\d+)(?:\.\d+)*)"
    r"\s*[.)\-–:]?\s+(?P<title>\S.*)$",
    re.IGNORECASE,
)
_HEADING_STYLE = re.compile(r"heading|đầu\s*đề|tiêu\s*đề", re.IGNORECASE)


def _level_of(num: str, seen_roman: bool) -> int:
    """Cấp của một đề mục suy từ chính con số của nó.

    Suy theo độ sâu con số là không đủ: "I." và "1." cùng một thành phần nên cùng
    ra cấp 1, khiến "1. Thông tin hệ thống" thay chỗ chương "I." thay vì nằm dưới.
    Trong mọi bản sizing thật, số LA MÃ là cấp chương và số Ả Rập là cấp dưới, nên
    chỉ cần biết tài liệu đã có chương La Mã hay chưa.
    """
    is_roman = bool(re.fullmatch(r"[IVXLCDM]+", num, re.IGNORECASE))
    base = 1 if is_roman or not seen_roman else 2
    return base + num.count(".")


@dataclass
class Element:
    """Một phần tử trong tài liệu, kèm đủ thông tin để neo một finding vào."""

    index: int                      # thứ tự xuất hiện, ổn định giữa các lần đọc
    kind: str                       # heading | paragraph | table | image
    text: str = ""                  # bảng: nội dung đã làm phẳng, để tìm kiếm
    page: int | None = None         # None = không suy được (xem docstring module)
    section: str = ""               # "IV.1.2" — số đề mục đang có hiệu lực
    section_title: str = ""         # "Định cỡ máy chủ BigData"
    level: int | None = None        # chỉ cho heading
    rows: list[list[str]] | None = None   # chỉ cho table
    style: str = ""

    @property
    def location(self) -> str:
        """Chuỗi neo dùng trong finding — cùng dạng người thẩm định hay viết."""
        parts = []
        if self.section:
            parts.append(f"Mục {self.section}")
        if self.page is not None:
            parts.append(f"trang {self.page}")
        return ", ".join(parts) or f"phần tử #{self.index}"


@dataclass
class DocxDocument:
    path: str
    elements: list[Element] = field(default_factory=list)
    page_source: str = "none"       # rendered | manual | none — độ tin của `page`
    warnings: list[str] = field(default_factory=list)

    @property
    def n_pages(self) -> int | None:
        pages = [e.page for e in self.elements if e.page is not None]
        return max(pages) if pages else None

    def tables(self) -> list[Element]:
        return [e for e in self.elements if e.kind == "table"]

    def images(self) -> list[Element]:
        return [e for e in self.elements if e.kind == "image"]

    def full_text(self) -> str:
        return "\n".join(e.text for e in self.elements if e.text)

    def by_section(self, section: str) -> list[Element]:
        """Phần tử thuộc một mục và mọi mục con của nó (IV.1 lấy cả IV.1.2)."""
        return [e for e in self.elements
                if e.section == section or e.section.startswith(section + ".")]


def _cell_text(tc) -> str:
    """Text của một ô, gồm cả bảng lồng bên trong."""
    return " ".join(
        "".join(n.text or "" for n in p.iter(qn("w:t"))).strip()
        for p in tc.iter(qn("w:p"))
    ).strip()


def _classify_heading(text: str, style: str, bold: bool,
                      seen_roman: bool = False) -> tuple[bool, str, str, int | None]:
    """(là đề mục?, số mục, tiêu đề, cấp).

    Style Heading là bằng chứng chắc nhất. Không có thì dựa vào mẫu đánh số — nhưng
    phải kèm điều kiện ngắn và (đậm | chữ hoa), nếu không mọi dòng bắt đầu bằng số
    trong bảng số liệu đều bị nhận nhầm thành đề mục.
    """
    by_style = bool(_HEADING_STYLE.search(style))
    m = _NUM.match(text)
    if by_style:
        lvl = None
        d = re.search(r"(\d+)", style)
        if d:
            lvl = int(d.group(1))
        if m:
            # Có bản gán "Heading 1" cho MỌI đề mục, kể cả mục con ("1. Thông tin
            # hệ thống" nằm dưới chương "I."). Khi text đã ghi rõ số thì con số
            # đáng tin hơn style, nên suy cấp từ nó.
            return True, m.group("num"), m.group("title").strip(), _level_of(m.group("num"), seen_roman)
        return True, "", text.strip(), lvl

    if m and len(text) <= 200:
        looks_like = bold or (text.strip() == text.strip().upper())
        if looks_like:
            num = m.group("num")
            # Suy cấp từ độ sâu con số là không đủ: "I." và "1." cùng có một thành
            # phần nên cùng ra cấp 1, khiến "1. Thông tin hệ thống" thay chỗ chương
            # "I." thay vì nằm dưới nó. Trong mọi bản sizing thật, số LA MÃ là cấp
            # chương và số Ả Rập là cấp dưới, nên chỉ cần biết tài liệu đã có chương
            # La Mã hay chưa là đặt đúng cấp.
            return True, num, m.group("title").strip(), _level_of(num, seen_roman)
    return False, "", "", None


def _compose(stack: dict[int, str], level: int, label: str) -> str:
    """Ghép số mục đầy đủ từ ngăn xếp theo cấp heading.

    Một quy tắc dùng chung cho cả số Word tự sinh lẫn số người gõ tay, vì cả hai
    đều chỉ hiện phần TƯƠNG ĐỐI: bản PNM hiện "4.1" ở cấp 3 (đã gồm cấp 2), bản
    Campaign hiện "1.4" ở cấp 2 (chưa gồm cấp 1). Nhãn nào đã chứa sẵn nhãn cha
    thì cắt phần trùng, còn lại nối thẳng — ra đúng dạng người thẩm định trích
    dẫn trong PNX ("Mục IV.1.5").
    """
    stack[level] = label
    for deeper in [k for k in stack if k > level]:
        del stack[deeper]

    parts: list[str] = []
    prev: str | None = None
    for lv in sorted(stack):
        piece = stack[lv]
        if prev and piece.startswith(prev + "."):
            piece = piece[len(prev) + 1:]
        parts.append(piece)
        prev = stack[lv]
    return ".".join(p for p in parts if p)


def read_docx(path: str) -> DocxDocument:
    doc = _Docx(path)
    out = DocxDocument(path=path)
    numbering = load_numbering(doc)

    # Ảnh: đếm trước để đối chiếu với số ảnh nhặt được khi duyệt thân bài.
    with zipfile.ZipFile(path) as z:
        n_media = sum(1 for n in z.namelist() if n.startswith("word/media/"))

    page = 1
    saw_rendered = saw_manual = False
    cur_num = cur_title = ""
    # Ngăn xếp số mục theo CẤP heading. Cần vì tài liệu thật dùng numId riêng cho
    # mỗi chương ở cấp 2, nên Word hiển thị "1." lặp lại dưới mọi chương — hai mục
    # khác nhau sẽ trùng `section`. Người thẩm định thì tự ghép ("Mục IV.1.1"),
    # nên ta ghép y như vậy: nối thành phần CUỐI của nhãn từng cấp.
    sec_stack: dict[int, str] = {}
    seen_roman = False
    idx = 0

    body = doc.element.body
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}

    for child in body.iterchildren():
        tag = child.tag

        if tag == qn("w:p"):
            # Ngắt trang phải đếm TRƯỚC khi gán trang cho đoạn này.
            if child.find(".//" + qn("w:lastRenderedPageBreak")) is not None:
                page += 1
                saw_rendered = True
            for br in child.iter(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    page += 1
                    saw_manual = True

            p = para_map.get(child)
            text = (p.text if p is not None else
                    "".join(n.text or "" for n in child.iter(qn("w:t")))).strip()
            style = (p.style.name if p is not None and p.style is not None else "")
            bold = any(bool(r.bold) for r in (p.runs if p is not None else []))

            has_img = (child.find(".//" + qn("w:drawing")) is not None
                       or child.find(".//" + qn("w:pict")) is not None)
            if has_img:
                out.elements.append(Element(
                    index=idx, kind="image", text=text, page=page,
                    section=cur_num, section_title=cur_title, style=style))
                idx += 1
                if not text:
                    continue

            if not text:
                continue

            is_h, num, title, lvl = _classify_heading(text, style, bold, seen_roman)
            if is_h and num and re.fullmatch(r"[IVXLCDM]+", num, re.IGNORECASE):
                seen_roman = True
            if is_h:
                # Word đánh số tự động: con số KHÔNG nằm trong text, phải dựng lại
                # từ numbering.xml, nếu không 28/48 bản sizing thật mất số mục.
                if not num:
                    ref = para_num_ref(child)
                    if ref is not None:
                        auto = numbering.label(*ref)
                        if auto:
                            if lvl is None:
                                lvl = ref[1] + 1
                            num, title = auto, text.strip()
                if num and lvl:
                    num = _compose(sec_stack, lvl, num)
                if num:
                    cur_num, cur_title = num, title
                else:
                    cur_title = title
                out.elements.append(Element(
                    index=idx, kind="heading", text=text, page=page,
                    section=cur_num, section_title=cur_title, level=lvl, style=style))
            else:
                out.elements.append(Element(
                    index=idx, kind="paragraph", text=text, page=page,
                    section=cur_num, section_title=cur_title, style=style))
            idx += 1

        elif tag == qn("w:tbl"):
            if child.find(".//" + qn("w:lastRenderedPageBreak")) is not None:
                page += 1
                saw_rendered = True
            rows: list[list[str]] = []
            for tr in child.findall(qn("w:tr")):
                rows.append([_cell_text(tc) for tc in tr.findall(qn("w:tc"))])
            flat = "\n".join(" | ".join(r) for r in rows)
            out.elements.append(Element(
                index=idx, kind="table", text=flat, page=page,
                section=cur_num, section_title=cur_title, rows=rows))
            idx += 1

    out.page_source = ("rendered" if saw_rendered else
                       "manual" if saw_manual else "none")
    if out.page_source == "none":
        out.warnings.append(
            "Không có dấu ngắt trang nào trong file — KHÔNG suy được số trang. "
            "Finding sẽ chỉ neo theo mục, không neo theo trang.")
        for e in out.elements:
            e.page = None
    elif out.page_source == "manual":
        out.warnings.append(
            "Chỉ có ngắt trang thủ công, không có dấu Word kết xuất — số trang là "
            "ƯỚC LƯỢNG, có thể lệch so với bản in.")

    if n_media and not out.images():
        out.warnings.append(
            f"File chứa {n_media} ảnh nhưng không nhặt được ảnh nào theo vị trí — "
            "C2 sẽ không có ngữ cảnh cho chúng.")
    if not any(e.section for e in out.elements):
        out.warnings.append(
            "Không nhận ra đề mục đánh số nào — finding chỉ neo được theo số thứ tự "
            "phần tử, rất khó đối chiếu với PNX.")
    return out
